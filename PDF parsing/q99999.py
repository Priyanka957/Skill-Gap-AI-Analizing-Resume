import streamlit as st
import re
import pdfplumber
import docx
import matplotlib.pyplot as plt
import plotly.graph_objects as go
from fpdf import FPDF
import pandas as pd
from datetime import datetime
import os
import plotly.express as px


from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity


# ========== SESSION STATE INITIALIZATION ==========
if 'ignore_swap_warning' not in st.session_state:
    st.session_state.ignore_swap_warning = False
if 'temp_resume' not in st.session_state:
    st.session_state.temp_resume = None
if 'temp_jd' not in st.session_state:
    st.session_state.temp_jd = None

# Check for swapped files in session state
if st.session_state.temp_resume and st.session_state.temp_jd:
    # Swap the files
    resume_file = st.session_state.temp_resume
    jd_file = st.session_state.temp_jd
    # Clear the temp storage
    st.session_state.temp_resume = None
    st.session_state.temp_jd = None
    st.session_state.ignore_swap_warning = True
    st.info("✅ Files have been swapped! Resume and Job Description are now in correct positions.")
st.markdown("""
    <style>
        .stProgress > div > div > div {
            background-color: var(--bar-color) !important;
        }
    </style>
""", unsafe_allow_html=True)

# ================= UTILS =================

def normalize_skill(skill):
    return skill.lower().strip()
SOFT_SKILL_SYNONYMS = {
    "communication": [
        "communication", "verbal communication", "written communication"
    ],
    "teamwork": ["teamwork", "team player", "collaboration"],
    "problem solving": ["problem solving", "analytical thinking"],
    "quick learner": ["quick learner", "fast learner", "self learner"],
    "hardworking": ["hardworking", "work ethic"],
    "time management": ["time management"],
    "adaptability": ["adaptability", "flexibility"],
}

def get_soft_skill_status(jd_skill, resume_skills):
    jd_norm = normalize_skill(jd_skill)

    # 1️⃣ Exact match
    for r in resume_skills:
        if normalize_skill(r) == jd_norm:
            return "exact"

    # 2️⃣ Synonym-based partial match ONLY
    if jd_norm in SOFT_SKILL_SYNONYMS:
        for r in resume_skills:
            if normalize_skill(r) in SOFT_SKILL_SYNONYMS[jd_norm]:
                return "partial"

    # 3️⃣ Otherwise missing
    return "missing"

SYNONYM_MAP = {
    "machine learning": [
        "machine learning",
        "basic machine learning concepts",
        "ml",
        "machine learning basics"
    ],
    "data visualization": [
        "data visualization",
        "visualization",
        "matplotlib",
        "power bi",
        "powerbi",
        "tableau",
        "data viz",
        "visualisation"
    ],
    "sql": [
        "sql",
        "database querying",
        "dbms",
        "database management",
        "sql programming"
    ],
    "python": [
        "python",
        "python programming",
        "python programmer",
        "numpy",
        "pandas",
        "matplotlib",
        "scikit-learn"
    ],
    "excel": [
        "excel",
        "microsoft excel",
        "ms excel",
        "advanced excel",
        "excal"  # fix your typo case
    ],
    "data analysis": [
        "data analysis",
        "data analytics",
        "data interpretation",
        "analyzing data",
        "data analyst"
    ]
}

def build_synonym_lookup():
    lookup = {}
    for canonical, variants in SYNONYM_MAP.items():
        for v in variants:
            lookup[v] = canonical
    return lookup

SYNONYM_LOOKUP = build_synonym_lookup()

# 🔥 REPLACE OLD FUNCTION WITH THIS
def get_match_status(jd_skill, resume_skills):
    jd_norm = normalize_skill(jd_skill)

    # 1️⃣ Exact match (STRICT)
    for r in resume_skills:
        if normalize_skill(r) == jd_norm:
            return "exact"

    # 2️⃣ Controlled partial match via synonyms ONLY
    for canonical, variants in SYNONYM_MAP.items():
        if jd_norm in variants:
            for r in resume_skills:
                r_norm = normalize_skill(r)
                if r_norm == canonical or r_norm in variants:
                    return "partial"

    # 3️⃣ Truly missing
    return "missing"
def detect_file_type_by_content(text):
    """Detect if text is from Resume or Job Description"""
    if not text or len(text.strip()) < 50:
        return "unknown"
    
    text_lower = text.lower()
    
    # Resume keywords
    resume_keywords = [
        'experience', 'education', 'skills', 'summary', 'objective',
        'certification', 'project', 'achievement', 'personal details',
        'phone', 'email', 'address', 'linkedin', 'github', 'worked',
        'developed', 'implemented', 'responsible for', 'university',
        'college', 'degree', 'bachelor', 'master', 'present'
    ]
    
    # Job Description keywords
    jd_keywords = [
        'job description', 'requirements', 'qualifications', 'responsibilities',
        'must have', 'should have', 'nice to have', 'required skills',
        'looking for', 'position overview', 'role', 'apply now',
        'submit your', 'send your', 'candidate should', 'we are seeking',
        'minimum requirements', 'preferred qualifications', 'essential',
        'benefits', 'salary', 'location', 'remote', 'hybrid'
    ]
    
    resume_score = sum(1 for keyword in resume_keywords if keyword in text_lower)
    jd_score = sum(1 for keyword in jd_keywords if keyword in text_lower)
    
    if resume_score > jd_score and resume_score >= 3:
        return "resume"
    elif jd_score > resume_score and jd_score >= 3:
        return "job_description"
    else:
        return "unknown"

def check_files_swapped(resume_text, jd_text):
    """Check if files might be swapped and return warning message"""
    resume_type = detect_file_type_by_content(resume_text)
    jd_type = detect_file_type_by_content(jd_text)
    
    if resume_type == "job_description" and jd_type == "resume":
        return True, """
        ⚠️ **WARNING: Files appear to be SWAPPED!**
        
        It looks like you uploaded:
        - **Job Description** in the **Resume** field
        - **Resume** in the **Job Description** field
        
        **Please upload the correct files in the correct fields!**
        """
    elif resume_type == "job_description":
        return True, """
        ⚠️ **WARNING: Possible file error!**
        
        The file uploaded as **Resume** appears to be a **Job Description**.
        Please check if you uploaded the correct file.
        """
    elif jd_type == "resume":
        return True, """
        ⚠️ **WARNING: Possible file error!**
        
        The file uploaded as **Job Description** appears to be a **Resume**.
        Please check if you uploaded the correct file.
        """
    
    return False, ""

def clean_text(text):
    """Clean extracted text from PDF/DOCX files while preserving structure"""
    if not text:
        return ""
    
    # Remove cid patterns (common in some PDFs)
    text = re.sub(r'\$cid[:]*\d+\$', '', text, flags=re.IGNORECASE)
    
    # Fix common PDF extraction issues
    text = re.sub(r'(\w)([A-Z])', r'\1 \2', text)  # Fix camelCase without spaces
    
    # Normalize whitespace (but preserve paragraph breaks)
    text = re.sub(r'[ \t]+', ' ', text)  # Multiple spaces/tabs to single space
    
    # Preserve important newlines but remove excessive ones
    text = re.sub(r'\n{4,}', '\n\n\n', text)  # 4+ newlines to 3
    text = re.sub(r'\n{3}', '\n\n', text)     # 3 newlines to 2
    
    # Remove trailing/leading whitespace from each line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join([line for line in lines if line])  # Remove empty lines
    
    return text.strip()

# ================= UTILS =================

def clean_text(text):
    """Clean extracted text from PDF/DOCX files - PRESERVE URLs"""
    # Remove cid patterns
    text = re.sub(r'\$cid[:]*\d+\$', '', text, flags=re.IGNORECASE)
    
    # Fix camelCase spacing
    text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
    
    # Normalize whitespace
    text = re.sub(r'[ \t]+', ' ', text)
    
    # Remove excessive newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()
def detect_sections(text):
    sections = {}
    current_section = "general"
    sections[current_section] = []

    for line in text.split("\n"):
        line_strip = line.strip().upper()

        # Detect headings
        if line_strip in [
            "TECHNICAL SKILLS",
            "SKILLS",
            "TECHNICAL SKILLS:",
            "EDUCATIONAL QUALIFICATION",
            "PROJECT",
            "CAREER OBJECTIVE"
        ]:
            current_section = line_strip.replace(":", "")
            sections[current_section] = []
        else:
            sections[current_section].append(line)

    return sections

def extract_text(file):
    """Extract ALL text from uploaded file - COMPLETE extraction"""
    if file is None:
        return ""
    
    # Save current position and reset to beginning
    try:
        original_position = file.tell()
        file.seek(0)
    except:
        pass  # Some file objects might not support seek
    
    try:
        if file.type == "application/pdf":
            text = ""
            try:
                with pdfplumber.open(file) as pdf:
                    for page in pdf.pages:
                        # METHOD 1: Standard extraction
                        page_text = page.extract_text()
                        
                        # METHOD 2: If standard fails, try layout extraction
                        if not page_text or len(page_text.strip()) < 20:
                            page_text = page.extract_text(layout=True)
                        
                        # METHOD 3: Try with looser parameters
                        if not page_text or len(page_text.strip()) < 20:
                            page_text = page.extract_text(
                                x_tolerance=3,
                                y_tolerance=3,
                                layout=False,
                                keep_blank_chars=False
                            )
                        
                        # METHOD 4: Last resort - try char-by-char extraction
                        if not page_text or len(page_text.strip()) < 10:
                            try:
                                page_text = page.extract_text(
                                    x_tolerance=5,
                                    y_tolerance=5,
                                    layout=True,
                                    keep_blank_chars=True
                                )
                            except:
                                pass
                        
                        if page_text:
                            text += page_text + "\n\n"  # Extra newline between pages
            except Exception as e:
                print(f"PDF extraction error: {e}")
                return ""
            
            # Reset file pointer
            try:
                file.seek(original_position)
            except:
                pass
            
            return text
        
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            try:
                doc = docx.Document(file)
                # Extract ALL text including tables
                all_text = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        all_text.append(para.text)
                
                # Also extract from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                all_text.append(cell.text)
                
                text = "\n".join(all_text)
                # Normalize bullets
                text = re.sub(r"[•▪●◦–—]", "-", text)
                
                # Reset file pointer
                try:
                    file.seek(original_position)
                except:
                    pass
                
                return text
            except Exception as e:
                print(f"DOCX extraction error: {e}")
                return ""
        
        elif file.type == "text/plain":
            try:
                content = file.read().decode("utf-8", errors='ignore')
                
                # Reset file pointer
                try:
                    file.seek(original_position)
                except:
                    pass
                
                return content
            except Exception as e:
                print(f"Text file reading error: {e}")
                return ""
        
        return ""
    
    except Exception as e:
        print(f"General extraction error: {e}")
        return ""

def extract_name_and_linkedin(text):
    """Extract candidate name and LinkedIn URL from resume text - returns tuple (name, linkedin_url)"""
    name = "Candidate"
    linkedin_url = ""
    
    # ===== 1. FIRST: Look for LinkedIn URL =====
    # Strategy 1: Look for Markdown format [text](url)
    markdown_pattern = r'\[([^\]]+)\]\(([^)]+)\)'
    markdown_matches = re.findall(markdown_pattern, text)
    
    for link_text, url in markdown_matches:
        if 'linkedin.com' in url.lower():
            linkedin_url = url
            break
        elif 'linkedin.com' in link_text.lower():
            linkedin_url = link_text
            break
    
    # Strategy 2: Look for "LinkedIn:" followed by URL
    if not linkedin_url:
        linkedin_label_pattern = r'LinkedIn[:]?\s*([^\s]+)'
        match = re.search(linkedin_label_pattern, text, re.IGNORECASE)
        if match:
            potential_url = match.group(1)
            if 'linkedin.com' in potential_url.lower():
                linkedin_url = potential_url
    
    # Strategy 3: Look for any LinkedIn URL in text
    if not linkedin_url:
        linkedin_url_patterns = [
            r'(https?://(?:www\.)?linkedin\.com/[^\s]+)',
            r'(www\.linkedin\.com/[^\s]+)',
            r'(linkedin\.com/[^\s]+)',
        ]
        
        for pattern in linkedin_url_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                linkedin_url = match.group(1)
                break
    
    # Clean the LinkedIn URL if found
    if linkedin_url:
        linkedin_url = linkedin_url.strip()
        # Remove any trailing punctuation
        linkedin_url = re.sub(r'[.,;:)\]]+$', '', linkedin_url)
        # Ensure proper protocol
        if not linkedin_url.startswith('http'):
            if linkedin_url.startswith('www.'):
                linkedin_url = 'https://' + linkedin_url
            elif linkedin_url.startswith('linkedin.com'):
                linkedin_url = 'https://www.' + linkedin_url
    
    # ===== 2. SECOND: Look for Name =====
    # Try to find name from the text
    lines = text.split('\n')
    
    # Pattern 1: Look for name in first few lines (most resumes have name at top)
    for line in lines[:10]:
        line = line.strip()
        if line:
            # Remove common prefixes/suffixes
            clean_line = re.sub(r'^[#\*\s]+|[#\*\s]+$', '', line)
            words = clean_line.split()
            
            # Check if it looks like a name (2-4 words, capitalized)
            if 2 <= len(words) <= 4:
                capitalized_words = sum(1 for w in words if w and w[0].isupper())
                if capitalized_words >= len(words) - 1:
                    # Make sure it's not a URL or other metadata
                    if not any(x in line.lower() for x in ['@', 'http', '://', 'www.', 'linkedin', 'phone', 'mobile']):
                        name = clean_line
                        break
    
    # Pattern 2: Look for "Name:" pattern
    if name == "Candidate":
        for line in lines:
            name_match = re.search(r'Name[:]?\s*([A-Z][a-zA-Z]*(?:\s+[A-Z][a-zA-Z]*)+)', line, re.IGNORECASE)
            if name_match:
                name = name_match.group(1).strip()
                break
    
    # Pattern 3: Extract from email if present
    if name == "Candidate":
        email_match = re.search(r'([a-zA-Z0-9._%+-]+)@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
        if email_match:
            email_local = email_match.group(1)
            if '.' in email_local:
                name_parts = email_local.split('.')
                name_parts = [p for p in name_parts if p and not p.isdigit()]
                if len(name_parts) >= 2:
                    name = ' '.join([p.title() for p in name_parts[:2]])
    
    return name, linkedin_url

def extract_professional_summary(text):
    """
    Extract the FULL professional summary exactly as written in the resume.
    Handles inline + multi-line summaries.
    """

    text = re.sub(r"\n{2,}", "\n", text)
    lines = text.split("\n")

    START_HEADERS = [
        "PROFESSIONAL SUMMARY",
        "SUMMARY",
        "CAREER OBJECTIVE",
        "OBJECTIVE",
        "PROFILE",
        "ABOUT ME"
    ]
    
    TOP_HEADERS = [
    "TECHNICAL SKILLS",
    "SOFT SKILLS",
    "SKILLS",
    "PROJECTS",
    "EDUCATION",
    "EXPERIENCE",
    "CERTIFICATIONS",
    "ACHIEVEMENTS"
]

    def detect_sections(text):
        sections = {}
        current_section = "GENERAL"
        sections[current_section] = []

        for line in text.split("\n"):
            line_strip = line.strip().upper().rstrip(":")

            if line_strip in TOP_HEADERS:
                current_section = line_strip
                sections[current_section] = []
            else:
                sections[current_section].append(line)

        return sections



    collected = []
    capturing = False

    for line in lines:
        clean = line.strip()
        upper = clean.upper()

        # Start capture
        for header in START_HEADERS:
            if upper.startswith(header):
                capturing = True

                # Handle inline summary
                inline = clean[len(header):].strip(" :-")
                if inline:
                    collected.append(inline)
                break
        else:
            if capturing:
                # Stop when next section starts
                if any(upper.startswith(h) for h in TOP_HEADERS):
                    break

                if clean:
                    collected.append(clean)

    summary = " ".join(collected).strip()

    if len(summary) >= 30:
        return summary

    return "Professional summary not clearly specified in the resume."

# ================= COMPREHENSIVE SKILL LISTS =================

TECHNICAL_SKILLS = [
    
    # Add these for Data Analyst JD
    "Data Analysis",
    "Data Interpretation",
    "Microsoft Excel",
    "MS Excel",
    "Advanced Excel",
    "Data Visualization",
    "Matplotlib",
    "Power BI",
    "Database Querying",
    "DBMS",
    
    # Core Programming & Embedded
    "Python", "Numpy", "Pandas", "Num Py", "Embedded C", "C", "C++", "Java", "JavaScript",
    "Embedded Systems", "Microcontroller", "Microcontroller Programming", "Machine Learning","Deep Learning",
    "Arduino", "Arduino Nano", "Arduino Uno", "Raspberry Pi", "ESP32", "ESP8266",

    # Specific Microcontrollers & Tools
    "8051 Microcontroller", "8051", "PIC Microcontroller", "AVR",
    "Keil", "Keil uVision", "MPLAB", "Proteus", "Proteus Simulation",

    #Machine Learing & Deep Learning
    "Machine Learing", "Deep Learing", "NLP", "Streamlit",
    
    #Front-end 
    "HTML", "CSS", "React", "Nodejs",
        
    # Protocols & Interfacing
    "I2C", "SPI", "UART", "GPIO", "ADC", "DAC",
    "Sensor Interfacing", "Actuator", "Relay", "LCD", "LED",
    "IR Sensor", "IR Sensors", "Ultrasonic Sensor", "LM35",

    # HDL & FPGA
    "Verilog HDL", "Verilog", "VHDL", "FPGA",

    # Data & Other Tools
    "SQL", "MySQL", "Excel", "MATLAB", "Simulink",
    "Internet of Things", "IoT", "GSM", "Bluetooth", "WiFi Module",

    # General Technical
    "Data Structures", "Algorithms", "OOP", "Git", "GitHub",
    "Circuit Design", "PCB Design", "Firmware Development",
    "RTOS", "Embedded Linux"
]

SOFT_SKILLS = [
    "Communication", "Verbal Communication", "Written Communication",
    "Teamwork", "Collaboration", "Leadership",
    "Problem Solving", "Critical Thinking", "Analytical Thinking",
    "Time Management", "Adaptability", "Flexibility",
    "Quick Learner", "Fast Learner", "Self Learner", "Eager to Learn",
    "Hardworking", "Work Ethic", "Dedication",
    "Active Listening", "Interpersonal Skills",
    "Presentation Skills", "Documentation"
]

def extract_additional_technical_terms(text):
    """
    Extract unknown technical terms (order preserved)
    Example: Arduino Nano, 8051, LM35, IR Sensor
    """
    results = []
    seen = set()

    lines = text.split("\n")
    for line in lines:
        # Only scan technical sections
        if any(k in line.lower() for k in ["technical", "skills", "tools", "technologies", "hardware"]):
            tokens = re.findall(r'\b[A-Z][A-Za-z0-9\-]{2,}\b', line)
            for t in tokens:
                key = t.lower()
                if key not in seen:
                    seen.add(key)
                    results.append(t)

    return results

def extract_skills(text):
    """
    Improved skill extraction:
    - Technical skills: searched in whole resume
    - Soft skills: searched ONLY after 'Soft Skills' heading (prevents leakage from degree names)
    - Basic false-positive filtering
    """
    skills = []
    seen = set()
    full_text_lower = text.lower()

    # ─── Try to isolate soft skills section to avoid false positives ───
    soft_section_text = full_text_lower
    if "soft skills" in full_text_lower:
        parts = full_text_lower.split("soft skills", 1)
        if len(parts) > 1:
            # Take text after "soft skills" until next major section or ~2000 chars
            remaining = parts[1]
            cutoff_phrases = ["education", "projects", "experience", "certification", "hobbies", "declaration"]
            for phrase in cutoff_phrases:
                if phrase in remaining:
                    remaining = remaining.split(phrase, 1)[0]
            soft_section_text = remaining[:2000]  # safety limit

    # 1. Technical skills – whole document
    for skill in TECHNICAL_SKILLS:
        pattern = r'\b' + re.escape(skill.lower()) + r'\b'
        if re.search(pattern, full_text_lower):
            key = skill.lower()
            if key not in seen:
                skills.append({"name": skill, "type": "Technical", "confidence": 92})
                seen.add(key)

    # 2. Soft skills – only in soft section (or fallback to whole text if no section found)
    for skill in SOFT_SKILLS:
        pattern = r'\b' + re.escape(skill.lower()) + r'\b'
        if re.search(pattern, soft_section_text):
            key = skill.lower()
            if key not in seen:
                skills.append({"name": skill, "type": "Soft", "confidence": 88})
                seen.add(key)

    # ─── Simple post-filters to remove common false positives ───
    # Post-filters - remove obvious false positives
    filtered_skills = []
    has_embedded_c = any("embedded c" in s["name"].lower() for s in skills)

    for s in skills:
        name_lower = s["name"].lower().strip()

        # Skip plain "C" when Embedded C exists
        if name_lower in ["c", "c programming", "c language"] and has_embedded_c:
            continue

        # Optional: skip very generic terms if you want
        if name_lower in ["sensor", "sensors"]:
            continue  # if you don't want generic "Sensor"

        filtered_skills.append(s)

    return filtered_skills

def normalize_and_merge_skills(skills_list):
    """
    Merge very similar / variant skill names into one canonical name.
    This prevents duplicates like 'Arduino' + 'Arduino Nano'.
    Expand the merge_rules dict as you see more real resumes.
    """
    # Define groups: "canonical_key": [possible variant substrings or exact names]
    merge_rules = {
        "arduino nano": ["arduino nano", "arduino", "nano", "arduino uno", "arduino-nano"],
        "8051 microcontroller": ["8051", "8051 microcontroller", "8051 micro controller"],
        "microcontroller": ["microcontroller", "microcontrollers", "mcu"],
        "ir sensor": ["ir sensor", "ir sensors", "ir blink sensor", "eye blink sensor", "ir eye sensor"],
        "microsoft excel": ["microsoft excel", "ms excel", "excel", "advanced excel"],
        "verilog hdl": ["verilog hdl", "verilog", "hdl", "veriloghdl"],
        "embedded c": ["embedded c", "embedded c programming", "c (embedded)"],
        # You can add more rules here later
    }

    canonical_map = {}

    for skill in skills_list:
        name_lower = skill["name"].lower().strip().replace("-", " ")

        found = False
        for canon_key, variants in merge_rules.items():
            if any(v in name_lower for v in variants) or name_lower == canon_key:
                display_name = canon_key.title()  # Use clean canonical name
                if canon_key not in canonical_map or skill["confidence"] > canonical_map[canon_key]["confidence"]:
                    canonical_map[canon_key] = {
                        "name": display_name,
                        "type": skill["type"],
                        "confidence": max(skill["confidence"], canonical_map.get(canon_key, {}).get("confidence", 0))
                    }
                found = True
                break

        if not found:
            key = name_lower
            if key not in canonical_map or skill["confidence"] > canonical_map[key]["confidence"]:
                canonical_map[key] = skill

    # Return sorted list: Technical first, then Soft
    merged = list(canonical_map.values())
    tech = sorted([s for s in merged if s["type"] == "Technical"], key=lambda x: x["name"].lower())
    soft = sorted([s for s in merged if s["type"] == "Soft"], key=lambda x: x["name"].lower())
    
    return tech + soft  # ← This is the final return — nothing after this!

def autopct_hide_zeros(pct):
    return f"{pct:.0f}%" if pct > 0 else ""


def skill_distribution_chart(tech_count, soft_count, title):
    labels = []
    values = []
    
    if tech_count > 0:
        labels.append("Technical Skills")
        values.append(tech_count)
    
    if soft_count > 0:
        labels.append("Soft Skills")
        values.append(soft_count)
    
    if not values:
        # Create an empty figure with message
        fig = go.Figure()
        fig.add_annotation(
            text="No Skills Found",
            xref="paper", yref="paper",
            x=0.5, y=0.5,
            showarrow=False,
            font=dict(size=20)
        )
        fig.update_layout(
            title=dict(text=title, font=dict(size=18)),
            showlegend=False,
            height=400,
            width=400
        )
        return fig
    
    # Create donut chart with fixed size
    fig = go.Figure(data=[go.Pie(
        labels=labels,
        values=values,
        hole=0.4,
        marker_colors=['#1f77b4', '#ff7f0e'],  # Consistent colors
        textinfo='percent+label',
        textposition='inside',
        textfont=dict(size=14, color='white'),
        hoverinfo='label+value+percent'
    )])
    
    fig.update_layout(
        title=dict(text=title, font=dict(size=18)),
        showlegend=True,
        legend=dict(
            orientation="h",
            yanchor="bottom",
            y=-0.1,
            xanchor="center",
            x=0.5
        ),
        height=400,  # Fixed height
        width=400,   # Fixed width
        margin=dict(t=50, b=50, l=50, r=50)  # Equal margins
    )
    
    return fig





def display_skill(skill):
    skill = skill.strip()
    lower = skill.lower()
    
    if lower == "sql":
        return "SQL"
    if "verilog" in lower:
        return "Verilog HDL"
    if "ir " in lower or "ir sensor" in lower:
        return "IR Sensor"
    if "embedded c" in lower:
        return "Embedded C"
    if lower == "c":
        return "C"  # But filter should remove it
    
    return skill.title()

# ================= MAIN APP =================

st.set_page_config(
    page_title="Skill Gap AI Analyzer",
    layout="wide"
)

# COMPLETE CSS (replace your current CSS block with this)
st.markdown("""
    <style>
    /* Remove all white space */
    .block-container {padding-top: 2rem; padding-bottom: 0;}
    .main .block-container {padding-top: 1rem;}
    
    /* Modern header */
    .main-header {
        background: linear-gradient(90deg, #1e3c72 0%, #2a5298 100%);
        padding: 25px 30px;
        border-radius: 0;
        margin: -1rem -1rem 20px -1rem;
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.1);
    }
    
    .main-header h1 {
        color: white;
        font-size: 2.2rem;
        margin: 0 0 8px 0;
        font-weight: 700;
        letter-spacing: -0.5px;
    }
    
    .main-header p {
        color: rgba(255, 255, 255, 0.85);
        margin: 0;
        font-size: 1rem;
        font-weight: 400;
    }
    
    /* Section headers - clean and minimal */
    .section-title {
        color: #1e3c72;
        font-size: 1.5rem;
        font-weight: 600;
        margin: 25px 0 15px 0;
        padding-bottom: 10px;
        border-bottom: 2px solid #1e3c72;
    }
    
    .sub-section {
        color: #2a5298;
        font-size: 1.2rem;
        font-weight: 600;
        margin: 20px 0 12px 0;
    }
    
    /* Cards - clean and compact */
    .info-card {
        background: #f8f9fa;
        border-left: 4px solid #1e3c72;
        padding: 15px;
        border-radius: 6px;
        margin: 15px 0;
    }
    
    .metric-card {
        background: white;
        padding: 15px;
        border-radius: 8px;
        border: 1px solid #e1e5e9;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        text-align: center;
    }
    
    .metric-value {
        font-size: 1.8rem;
        font-weight: 700;
        color: #1e3c72;
        margin: 0;
    }
    
    .metric-label {
        font-size: 0.9rem;
        color: #6c757d;
        margin: 5px 0 0 0;
    }
    
    /* Skills styling */
    .skills-container {
        background: white;
        padding: 15px;
        border-radius: 8px;
        border: 1px solid #e1e5e9;
        margin: 10px 0;
    }
    
    .skill-item {
        display: inline-block;
        background: #e8f0fe;
        color: #1e3c72;
        padding: 6px 12px;
        border-radius: 16px;
        margin: 4px;
        font-size: 0.9rem;
        font-weight: 500;
    }
    
    .skill-match { background: #d4edda; color: #155724; }
    .skill-partial { background: #fff3cd; color: #856404; }
    .skill-missing { background: #f8d7da; color: #721c24; }
    
    /* Status messages */
    .status-success {
        background: #d4edda;
        color: #155724;
        padding: 12px 15px;
        border-radius: 6px;
        border-left: 4px solid #28a745;
        margin: 15px 0;
        font-weight: 500;
    }
    
    .status-warning {
        background: #fff3cd;
        color: #856404;
        padding: 12px 15px;
        border-radius: 6px;
        border-left: 4px solid #ffc107;
        margin: 15px 0;
        font-weight: 500;
    }
    
    .status-error {
        background: #f8d7da;
        color: #721c24;
        padding: 12px 15px;
        border-radius: 6px;
        border-left: 4px solid #dc3545;
        margin: 15px 0;
        font-weight: 500;
    }
    
    /* Button styling */
    .stButton button {
        background: linear-gradient(90deg, #1e3c72 0%, #2a5298 100%);
        color: white;
        border: none;
        border-radius: 6px;
        font-weight: 500;
        transition: all 0.2s;
    }
    
    .stButton button:hover {
        background: linear-gradient(90deg, #2a5298 0%, #1e3c72 100%);
        transform: translateY(-1px);
        box-shadow: 0 4px 8px rgba(0,0,0,0.1);
    }
    
    /* Download buttons */
    .download-btn {
        background: linear-gradient(90deg, #28a745 0%, #20c997 100%);
        color: white;
        border: none;
        border-radius: 6px;
        padding: 10px 20px;
        font-weight: 500;
        width: 100%;
    }
    
    /* Remove spacing from streamlit elements */
    div[data-testid="stExpander"] { margin: 10px 0; }
    .streamlit-expanderHeader { padding: 12px 15px; }
    div[data-testid="stVerticalBlock"] { gap: 0.5rem; }
    
    /* File upload area */
    .upload-section {
        background: #f8f9fa;
        padding: 15px;
        border-radius: 8px;
        border: 2px dashed #ced4da;
        margin: 10px 0;
    }
    
    /* Table styling */
    .dataframe {
        border: 1px solid #e1e5e9;
        border-radius: 6px;
        margin: 10px 0;
    }
    
    /* Hide streamlit branding */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    
    /* Chart containers */
    .plot-container { margin: 5px 0; }
    
    /* Compact form elements */
    .stTextArea textarea { min-height: 150px; }
    .stSelectbox, .stRadio > div { margin: 5px 0; }
    
       /* Skill Similarity Scores Section */
    .similarity-section {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        padding: 20px;
        border-radius: 12px;
        border-left: 6px solid #4A90E2;
        margin: 20px 0;
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08);
    }
    
    .similarity-title {
        color: #2c3e50;
        font-size: 1.4rem;
        font-weight: 700;
        margin-bottom: 15px;
        padding-bottom: 10px;
        border-bottom: 2px solid #4A90E2;
    }
    
    /* Skill Priority Radar Section */
    .priority-section {
        background: linear-gradient(135deg, #fffde7 0%, #fff9c4 100%);
        padding: 20px;
        border-radius: 12px;
        border-left: 6px solid #FFA726;
        margin: 20px 0;
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.08);
    }
    
    .priority-title {
        color: #5D4037;
        font-size: 1.4rem;
        font-weight: 700;
        margin-bottom: 15px;
        padding-bottom: 10px;
        border-bottom: 2px solid #FFA726;
    }
    
    /* High-Impact Skills Section */
    .high-impact-section {
        background: linear-gradient(135deg, #e8f5e9 0%, #c8e6c9 100%);
        padding: 15px;
        border-radius: 10px;
        margin: 15px 0;
        border: 1px solid #81C784;
    }
    
    .high-impact-title {
        color: #1B5E20;
        font-size: 1.2rem;
        font-weight: 600;
        margin-bottom: 10px;
    }
    
    /* Skill list items */
    .skill-list-item {
        background: rgba(255, 255, 255, 0.9);
        padding: 10px 15px;
        margin: 8px 0;
        border-radius: 8px;
        border-left: 4px solid #3498db;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    
    .missing-skill-item {
        background: rgba(255, 243, 224, 0.9);
        padding: 10px 15px;
        margin: 8px 0;
        border-radius: 8px;
        border-left: 4px solid #FF9800;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    
    /* Dataframe styling */
    .dataframe {
        border: 2px solid #4CAF50 !important;
        border-radius: 10px !important;
        overflow: hidden !important;
        box-shadow: 0 4px 8px rgba(76, 175, 80, 0.1) !important;
    }
    
    .dataframe th {
        background: linear-gradient(90deg, #4CAF50, #45a049) !important;
        color: white !important;
        font-weight: bold !important;
        padding: 12px !important;
    }
    
    .dataframe td {
        padding: 10px !important;
        border-bottom: 1px solid #e0e0e0 !important;
    }
    
    .dataframe tr:nth-child(even) {
        background-color: #f9f9f9 !important;
    }
    
    .dataframe tr:hover {
        background-color: #e8f5e8 !important;
    }
    
    /* Info box styling */
    .colorful-info-box {
        background: linear-gradient(135deg, #e3f2fd 0%, #bbdefb 100%);
        padding: 15px;
        border-radius: 10px;
        border-left: 5px solid #2196F3;
        margin: 15px 0;
        font-style: italic;
    }
    
    /* Start with skill highlight */
    .start-with-box {
        background: linear-gradient(135deg, #FFF3E0 0%, #FFE0B2 100%);
        padding: 15px;
        border-radius: 10px;
        border: 2px solid #FF9800;
        margin: 15px 0;
        font-weight: 500;
    }
    
    /* Caption styling */
    .colorful-caption {
        color: #666;
        font-size: 0.9rem;
        font-style: italic;
        padding: 8px 0;
        background: rgba(245, 245, 245, 0.8);
        border-radius: 6px;
        padding: 10px;
        margin: 10px 0;
    }
    
    </style>
""", unsafe_allow_html=True)
    

# KEEP THIS HEADER (it's already correct)
st.markdown("""
    <div class="main-header">
        <h1>Skill Gap AI Analyzer</h1>
        <p>Data Ingestion · Skill Extraction · Gap Analysis · Dashboard & Reports</p>
    </div>
""", unsafe_allow_html=True)


# Add this RIGHT AFTER your main header but BEFORE "File Upload"
with st.expander("📋 **Upload Instructions**", expanded=True):
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("""
        **📄 Upload Resume (Your CV):**
        - Should contain your skills, experience, education
        - Usually starts with your name and contact info
        - Has sections like: Experience, Education, Skills
        """)
    
    with col2:
        st.markdown("""
        **📋 Upload Job Description:**
        - Should contain required skills and qualifications
        - Usually starts with job title and company
        - Has sections like: Requirements, Responsibilities
        """)
    
    st.warning("⚠️ **Important:** Don't mix them up! Resume in first field, Job Description in second field.")
# File Upload
col1, col2 = st.columns(2)

with col1:
    st.markdown('<div class="upload-section">', unsafe_allow_html=True)
    st.markdown("**Upload Resume**")
    resume_file = st.file_uploader(
        "Choose a file",
        type=["pdf", "docx", "txt"],
        key="resume",
        label_visibility="collapsed"
    )
    if resume_file:
        st.markdown(f'<div class="status-success">✓ {resume_file.name}</div>', unsafe_allow_html=True)
    else:
        st.markdown('<div style="color: #6c757d; font-size: 0.9rem; margin-top: 10px;">PDF, DOCX, or TXT files</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

with col2:
    st.markdown('<div class="upload-section">', unsafe_allow_html=True)
    st.markdown("**Upload Job Description**")
    jd_file = st.file_uploader(
        "Choose a file",
        type=["pdf", "docx", "txt"],
        key="jd",
        label_visibility="collapsed"
    )
    if jd_file:
        st.markdown(f'<div class="status-success">✓ {jd_file.name}</div>', unsafe_allow_html=True)
    else:
        st.markdown('<div style="color: #6c757d; font-size: 0.9rem; margin-top: 10px;">PDF, DOCX, or TXT files</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)
        
# Check if files are uploaded
if not resume_file or not jd_file:
    if resume_file and not jd_file:
        st.error("❌ **Please upload both Resume AND Job Description files.**")
        st.info("📝 You've uploaded a Resume but need to upload a Job Description too.")
    elif not resume_file and jd_file:
        st.error("❌ **Please upload both Resume AND Job Description files.**")
        st.info("📝 You've uploaded a Job Description but need to upload a Resume too.")
    else:
        st.info("👈 **Please upload both a Resume and a Job Description to start the analysis.**")
    
    # Add quick upload buttons for demo/testing
    with st.expander("🔄 Need test files?"):
        st.markdown("""
        **For testing purposes, you can:**
        1. Create a simple text file for Resume with:
        ```
        John Doe
        Software Engineer
        
        SKILLS
        Python, JavaScript, SQL
        
        EXPERIENCE
        Software Developer at ABC Corp (2020-Present)
        ```
        
        2. Create a simple text file for Job Description with:
        ```
        Senior Software Engineer
        
        REQUIREMENTS
        - 5+ years experience in Python
        - Strong knowledge of JavaScript
        - SQL database experience
        ```
        """)
    st.stop()



if resume_file and jd_file:
    # ========== FILE SWAP DETECTION ==========
    st.markdown('<div class="section-title">DATA INGESTION & PARSING</div>', unsafe_allow_html=True)
    
    # First extract a small amount of text to check file types
    resume_text_temp = ""
    jd_text_temp = ""
    
    # Save file positions first
    resume_pos = resume_file.tell()
    jd_pos = jd_file.tell()
    
    try:
        # Extract a small sample from each file
        resume_text_temp = extract_text(resume_file)
        resume_file.seek(resume_pos)  # Reset file pointer
        resume_text_temp = clean_text(resume_text_temp)
    except:
        resume_file.seek(resume_pos)  # Reset file pointer
    
    try:
        jd_text_temp = extract_text(jd_file)
        jd_file.seek(jd_pos)  # Reset file pointer
        jd_text_temp = clean_text(jd_text_temp)
    except:
        jd_file.seek(jd_pos)  # Reset file pointer
    
    # Check if files might be swapped
    if resume_text_temp and jd_text_temp:
        files_swapped, warning_message = check_files_swapped(resume_text_temp, jd_text_temp)
        
        if files_swapped:
            st.markdown(f'<div class="file-swap-warning">{warning_message}</div>', unsafe_allow_html=True)
            
            # Ask user what to do
            col1, col2 = st.columns(2)
            with col1:
                if st.button("🔄 Re-upload Correct Files", use_container_width=True, type="primary"):
                    # Clear the files
                    st.session_state.resume = None
                    st.session_state.jd = None
                    st.rerun()
            
            with col2:
                if st.button("⚠️ Continue Anyway", use_container_width=True, type="secondary"):
                    st.warning("Continuing with possible incorrect files...")
                    # Store that we warned the user
                    st.session_state.ignored_warning = True
            st.stop()  # Stop further processing
    
    # ========== END FILE SWAP DETECTION ==========
    
    # ────────────────────────────────────────────────────────────────
    # SAFE FILE READING WITH ERROR HANDLING
    # ────────────────────────────────────────────────────────────────
    resume_text = ""
    jd_text = ""

    # Read Resume
    try:
        resume_text = extract_text(resume_file)
        resume_text = clean_text(resume_text)
    except Exception as e:
        st.error(f"Failed to read resume file: {str(e)}")
        resume_text = ""

    # Read Job Description
    try:
        jd_text = extract_text(jd_file)
        jd_text = clean_text(jd_text)
    except Exception as e:
        st.error(f"Failed to read job description file: {str(e)}")
        jd_text = ""

    # User-friendly feedback if reading failed
    if not resume_text and not jd_text:
        st.markdown('<div class="file-error-box">⚠️ Could not read content from either file. Please check file format and try again.</div>', unsafe_allow_html=True)
        st.info("**Supported formats:** PDF, DOCX, TXT files with readable text (not scanned images)")
        st.stop()
    elif not resume_text:
        st.markdown('<div class="file-error-box">⚠️ Could not read the resume file. Please upload a valid resume document.</div>', unsafe_allow_html=True)
    elif not jd_text:
        st.markdown('<div class="file-error-box">⚠️ Could not read the job description file. Please upload a valid job description.</div>', unsafe_allow_html=True)
    
    # Check for minimum content
    if resume_text and len(resume_text.strip()) < 50:
        st.warning("⚠️ Resume appears very short. Minimum 50 characters required for proper analysis.")
    
    if jd_text and len(jd_text.strip()) < 50:
        st.warning("⚠️ Job Description appears very short. Minimum 50 characters required for proper analysis.")

    # ────────────────────────────────────────────────────────────────
    # Show previews only if we have at least some content
    # ────────────────────────────────────────────────────────────────
# ────────────────────────────────────────────────────────────────
# Show COMPLETE previews - NO LIMITS
# ────────────────────────────────────────────────────────────────
if resume_text or jd_text:
    col1, col2 = st.columns(2)

    with col1:
        st.markdown("### 📄 Resume Preview (COMPLETE CONTENT)")
        if resume_text:
            # Show ALL content - NO LIMIT
            st.text_area(
                "Resume Full Content",
                resume_text,  # ← NO [:2000] limit here!
                height=600,   # Taller to show more
                label_visibility="collapsed",
                key="resume_preview_full"
            )
            # Show exact character count
            st.success(f"✅ **FULL EXTRACTION COMPLETE:** {len(resume_text)} characters")
            
            # Optional: Add a download button for the extracted text
            st.download_button(
                label="📥 Download Extracted Resume Text",
                data=resume_text,
                file_name="extracted_resume.txt",
                mime="text/plain",
                use_container_width=True
            )
        else:
            st.info("Resume content could not be loaded.")

    with col2:
        st.markdown("### 📋 Job Description Preview (COMPLETE CONTENT)")
        if jd_text:
            # Show ALL content - NO LIMIT
            st.text_area(
                "Job Description Full Content",
                jd_text,  # ← NO [:2000] limit here!
                height=600,
                label_visibility="collapsed",
                key="jd_preview_full"
            )
            st.success(f"✅ **FULL EXTRACTION COMPLETE:** {len(jd_text)} characters")
            
            # Optional: Add a download button
            st.download_button(
                label="📥 Download Extracted Job Description Text",
                data=jd_text,
                file_name="extracted_job_description.txt",
                mime="text/plain",
                use_container_width=True
            )
        else:
            st.info("Job Description content could not be loaded.")
    
    # Check if we can proceed to Milestone 2
    if not resume_text.strip() or not jd_text.strip():
        st.error("""
        ❌ **Cannot proceed to analysis**
        
        **Reason:** One or both files are empty or unreadable.
        
        **Please:**
        1. Check if files are in correct format (PDF, DOCX, TXT)
        2. Ensure files contain readable text (not scanned images)
        3. Upload valid documents with sufficient content
        """)
        st.stop()
        
    # ================= DATA INGESTION COMPLETION MESSAGE =================
    st.success("✅ **Completed: Documents uploaded, parsed, cleaned and previewed successfully.**")

    # Milestone 2: Skill Extraction
   
    # Milestone 2: Skill Extraction using NLP
    resume_skills_raw = extract_skills(resume_text)
    resume_skills = normalize_and_merge_skills(resume_skills_raw)

    jd_skills_raw = extract_skills(jd_text)
    jd_skills = normalize_and_merge_skills(jd_skills_raw)  # good practice
    
    # ================= JD vs Resume Skill Alignment (Milestone 2) =================
    resume_skill_set = {s["name"].lower() for s in resume_skills}
    jd_skill_set = {s["name"].lower() for s in jd_skills}


    jd_matched = []
    jd_missing = []

    for skill in jd_skill_set:
        if skill in resume_skill_set:
            jd_matched.append(skill)
        else:
            jd_missing.append(skill)

    # ===== Skill-level similarity score calculation =====
    skill_scores = {}

    for skill in jd_skills:
        tfidf = TfidfVectorizer()
        vectors = tfidf.fit_transform([skill["name"], resume_text])
        score = cosine_similarity(vectors[0], vectors[1])[0][0]
        skill_scores[skill["name"]] = round(score, 2)

    # ===== Skill categorization =====
    tech_resume = [s for s in resume_skills if s["type"] == "Technical"]
    soft_resume = [s for s in resume_skills if s["type"] == "Soft"]
    tech_jd = [s for s in jd_skills if s["type"] == "Technical"]
    soft_jd = [s for s in jd_skills if s["type"] == "Soft"]

    # ================= Milestone 2: Skill Extraction =================
    st.markdown('<div class="section-title">SKILL EXTRACTION USING NLP</div>', unsafe_allow_html=True)

    # -------- Skill Extraction --------
    resume_skills_raw = extract_skills(resume_text)
    resume_skills = normalize_and_merge_skills(resume_skills_raw)

    jd_skills_raw = extract_skills(jd_text)
    jd_skills = normalize_and_merge_skills(jd_skills_raw)

    # -------- Candidate Info --------
    candidate_name, linkedin_url = extract_name_and_linkedin(resume_text)
    professional_summary = extract_professional_summary(resume_text)

    # -------- Categorize Skills --------
    tech_resume = [s for s in resume_skills if s["type"] == "Technical"]
    soft_resume = [s for s in resume_skills if s["type"] == "Soft"]
    tech_jd = [s for s in jd_skills if s["type"] == "Technical"]
    soft_jd = [s for s in jd_skills if s["type"] == "Soft"]

    # -------- Candidate Information --------
    st.markdown("##  Candidate Information")
    st.markdown(f"**Name:** {candidate_name}")

    if linkedin_url:
        st.markdown(f"**LinkedIn:** {linkedin_url}")
    else:
        st.info("LinkedIn not found")

    st.markdown(f"** Professional Summary:** {professional_summary}")

# -------- Extracted Skills --------
    st.markdown("---")
    st.markdown("##  Extracted Skills")

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("### Resume Skills")

        with st.expander("Technical Skills", expanded=True):
            st.write(", ".join([s["name"] for s in tech_resume]) or "None")

        with st.expander("Soft Skills", expanded=True):
            st.write(", ".join([s["name"] for s in soft_resume]) or "None")

    with col2:
        st.markdown("### Job Description Skills")

        with st.expander("Technical Skills", expanded=True):
            st.write(", ".join([s["name"] for s in tech_jd]) or "None")

        with st.expander("Soft Skills", expanded=True):
            st.write(", ".join([s["name"] for s in soft_jd]) or "None")
            
  


# ================= Skill Distribution =================
    st.markdown("---")
    st.markdown("##  Skill Distribution")

    c1, c2 = st.columns(2, gap="large")

    # -------- Resume --------
    with c1:
        fig_resume = skill_distribution_chart(
            len(tech_resume),
            len(soft_resume),
            "Resume Skill Distribution"
        )
        st.plotly_chart(fig_resume, use_container_width=False)  # Changed to plotly

        st.markdown("####  Resume Metrics")
        r1, r2, r3, r4 = st.columns(4)

        r1.metric("Technical Skills", len(tech_resume))
        r2.metric("Soft Skills", len(soft_resume))
        r3.metric("Total Skills", len(resume_skills))

        resume_avg_conf = (
            round(sum(s["confidence"] for s in resume_skills) / len(resume_skills), 1)
            if resume_skills else 0
        )
        r4.metric("Avg Confidence", f"{resume_avg_conf}%")

    # -------- Job Description --------
    with c2:
        fig_jd = skill_distribution_chart(
            len(tech_jd),
            len(soft_jd),
            "JD Skill Distribution"
        )
        st.plotly_chart(fig_jd, use_container_width=False)  # Changed to plotly

        st.markdown("####  JD Metrics")
        j1, j2, j3, j4 = st.columns(4)

        j1.metric("Technical Skills", len(tech_jd))
        j2.metric("Soft Skills", len(soft_jd))
        j3.metric("Total Skills", len(jd_skills))

        jd_avg_conf = (
            round(sum(s["confidence"] for s in jd_skills) / len(jd_skills), 1)
            if jd_skills else 0
        )
        j4.metric("Avg Confidence", f"{jd_avg_conf}%")

    # ================= SKILL EXTRACTION COMPLETION MESSAGE =================
    st.success("✅ **Completed: Skills extracted using NLP, candidate information identified successfully.**")


    # Milestone 3: Skill Gap Analysis
    st.markdown('<div class="section-title">SKILL GAP ANALYSIS & SIMILARITY MATCHING</div>', unsafe_allow_html=True)
    resume_skill_names = {s["name"].lower() for s in resume_skills}
    resume_skill_map = {
    normalize_skill(s["name"]): s["name"]
    for s in resume_skills
    }
    jd_skill_names = {s["name"].lower() for s in jd_skills}

    matched = set()
    partial = set()
    missing = set()

    for jd_skill in jd_skill_names:
        status = get_match_status(jd_skill, resume_skill_names)

        if status == "exact":
            matched.add(jd_skill)
        elif status == "partial":
            partial.add(jd_skill)
        else:
            missing.add(jd_skill)

    missing = jd_skill_names - matched - partial


    overall_match = int((len(matched) / len(jd_skill_names)) * 100) if jd_skill_names else 0

    left, right = st.columns([3, 2])

    with left:
        st.markdown("### Similarity Matrix")
        jd_list = sorted(jd_skill_names)
        resume_list = sorted([display_skill(s) for s in resume_skill_names])

        st.caption(
            f"Comparing {len(resume_list)} resume skills with {len(jd_list)} job description skills"
        )
        
        if resume_list and jd_list:
            # Limit number of skills for better visualization
            max_skills = 15
            if len(resume_list) > max_skills:
                jd_list_display = jd_list
                st.info(f"Showing first {max_skills} resume skills")
            else:
                resume_list_display = resume_list
                
            if len(jd_list) > max_skills:
                jd_list_display = jd_list[:max_skills]
                st.info(f"Showing first {max_skills} job skills")
            else:
                jd_list_display = jd_list

# ================= Similarity Matrix + Overview =================

                if jd_list_display:

                    # Create similarity matrix
                    fig = go.Figure()

                    STATUS_ROW = {
                        "exact": "✅ Exact Match",
                        "partial": "🟡 Partial Match",
                        "missing": "❌ Missing"
                    }

                    COLOR_MAP = {
                        "exact": "green",
                        "partial": "orange",
                        "missing": "red"
                    }

                    SIZE_MAP = {
                        "exact": 26,
                        "partial": 22,
                        "missing": 16
                    }

                    for jd_skill in jd_list_display:
                        jd_norm = normalize_skill(jd_skill)
                        status = "missing"
                        matched_resume_skill = None  # IMPORTANT

                        if jd_skill in SOFT_SKILLS:
                            status = get_soft_skill_status(jd_skill, resume_skill_names)
                        else:
                            status = get_match_status(jd_skill, resume_skill_names)

                        # Exact match → same skill
                        if status == "exact":
                            matched_resume_skill = display_skill(jd_skill)

                        # Partial match → controlled synonym match
                        elif status == "partial":
                            jd_norm = normalize_skill(jd_skill)

                            for canonical, variants in SYNONYM_MAP.items():
                                if jd_norm in variants:
                                    for r in resume_skill_names:
                                        if normalize_skill(r) == canonical or normalize_skill(r) in variants:
                                            matched_resume_skill = resume_skill_map.get(
                                                normalize_skill(r), r
                                            )
                                            break

                        fig.add_trace(go.Scatter(
                            x=[display_skill(jd_skill)],
                            y=[STATUS_ROW[status]],
                            mode="markers+text" if status == "partial" else "markers",
                            text=[matched_resume_skill] if status == "partial" else None,
                            textposition="top center",
                            marker=dict(
                                size=SIZE_MAP[status],
                                color=COLOR_MAP[status],
                                line=dict(color="black", width=1)
                            ),
                            hovertemplate=(
                                "<b>JD Skill:</b> %{x}<br>"
                                "<b>Status:</b> " + STATUS_ROW[status] + "<br>" +
                                "<b>Matched Resume Skill:</b> " +
                                (matched_resume_skill if matched_resume_skill else "—") +
                                "<extra></extra>"
                            ),
                            showlegend=False
                        ))

                    fig.update_layout(
                        height=380,
                        plot_bgcolor="white",
                        margin=dict(l=120, r=40, t=30, b=80),
                        xaxis=dict(
                            title="<b>Job Description Skills</b>",
                            tickangle=45,
                            showgrid=True,
                            gridcolor="lightgray"
                        ),
                        yaxis=dict(
                            title="<b>Match Status</b>",
                            categoryorder="array",
                            categoryarray=[
                                "✅ Exact Match",
                                "🟡 Partial Match",
                                "❌ Missing"
                            ],
                            showgrid=True,
                            gridcolor="lightgray"
                        ),
                        showlegend=False
                    )

                    st.plotly_chart(fig, use_container_width=True)
                st.markdown("### Missing Skills")
                if missing:
                    for skill in sorted(missing):
                        st.markdown(f"❌ **{skill.title()}**")
                else:
                    st.success("No missing skills! ✅")

            # ================= RIGHT SIDE : SKILL MATCH OVERVIEW =================

            with right:
                st.markdown("### Skill Match Overview")

                total_jd = len(jd_list_display)
                overall_match = int((len(matched) / total_jd) * 100) if total_jd else 0

                match_color = (
                    "#10B981" if overall_match >= 70
                    else "#F59E0B" if overall_match >= 50
                    else "#EF4444"
                )

                c1, c2 = st.columns(2)
                c3, c4 = st.columns(2)

                c1.markdown(f"""
                <div style="background:{match_color};padding:15px;border-radius:10px;text-align:center;">
                    <h3 style="color:white;margin:0;">{overall_match}%</h3>
                    <p style="color:white;margin:0;">Overall Match</p>
                </div>
                """, unsafe_allow_html=True)

                c2.metric("Matched Skills", len(matched))
                c3.metric("Partial Matches", len(partial))
                c4.metric("Missing Skills", len(missing))

                # --- Donut chart ---
                if matched or partial or missing:
                    total = len(matched) + len(partial) + len(missing)

                    # Create the donut chart
                    donut = go.Figure(go.Pie(
                        labels=["Matched", "Partial", "Missing"],
                        values=[len(matched), len(partial), len(missing)],
                        hole=0.65,
                        marker_colors=["#10B981", "#F59E0B", "#EF4444"],

                        # ✅ SHOW PERCENT ONLY ONCE
                        textinfo="percent",
                        textposition="inside",
                        textfont=dict(size=14, color="white"),
                        insidetextorientation="radial",

                        # Hover info (does NOT duplicate text)
                        hoverinfo="label+value"
                    ))

                    # ✅ Center annotation (overall match)
                    donut.add_annotation(
                        text=f"<b>{overall_match}%</b>",
                        x=0.5, y=0.5,
                        font=dict(size=28, color="#333", family="Arial Black"),
                        showarrow=False
                    )

                    donut.update_layout(
                        height=300,
                        margin=dict(t=20, b=20, l=20, r=20),
                        showlegend=True,
                        legend=dict(
                            orientation="h",
                            x=0.5,
                            xanchor="center",
                            y=-0.1,
                            font=dict(size=12)
                        ),
                        title=dict(
                            text="Skill Match Distribution",
                            x=0.5,
                            xanchor="center",
                            font=dict(size=16)
                        )
                    )

                    st.plotly_chart(donut, use_container_width=True)
                    
    # ================= SKILL GAP ANALYSIS COMPLETION MESSAGE =================
    st.success("✅ **Completed: Skill gap analysis and similarity matching performed successfully.**")


    # Milestone 4: Dashboard & Reports
    st.markdown('<div class="section-title">DASHBOARD & REPORT EXPORT</div>', unsafe_allow_html=True)

    # Prepare data for visualization
    all_skills = sorted(list(jd_skill_names | resume_skill_names))
    
    if all_skills:
        resume_scores = []
        jd_scores = []

        resume_set = {normalize_skill(s) for s in resume_skill_names}
        jd_set    = {normalize_skill(s) for s in jd_skill_names}

        for skill in all_skills:
            norm = normalize_skill(skill)
            
            # Resume coverage: high if present in resume
            resume_scores.append(95 if norm in resume_set else 10)
            
            # JD requirement: high if present in JD
            jd_scores.append(95 if norm in jd_set else 10)

        df_skills = pd.DataFrame({
            "Skill": [s.title() for s in all_skills],
            "Resume Skill %": resume_scores,
            "Job Requirement %": jd_scores
        })

        # Key metrics cards
        st.markdown("###  Performance Metrics")
        c1, c2, c3, c4 = st.columns(4)
        
        c1.markdown(f"""
        <div style="background:#E0F2FE;padding:20px;border-radius:10px;text-align:center;border-left:5px solid #0EA5E9;">
        <h2 style="color:#0EA5E9;margin:0;">{overall_match}%</h2>
        <p style="margin:0;">Overall Match</p>
        </div>
        """, unsafe_allow_html=True)
        
        c2.markdown(f"""
        <div style="background:#DCFCE7;padding:20px;border-radius:10px;text-align:center;border-left:5px solid #22C55E;">
        <h2 style="color:#166534;margin:0;">{len(matched)}</h2>
        <p style="margin:0;">Matched Skills</p>
        </div>
        """, unsafe_allow_html=True)
        
        c3.markdown(f"""
        <div style="background:#FEF3C7;padding:20px;border-radius:10px;text-align:center;border-left:5px solid #F59E0B;">
        <h2 style="color:#92400E;margin:0;">{len(partial)}</h2>
        <p style="margin:0;">Partial Matches</p>
        </div>
        """, unsafe_allow_html=True)
        
        c4.markdown(f"""
        <div style="background:#FEE2E2;padding:20px;border-radius:10px;text-align:center;border-left:5px solid #EF4444;">
        <h2 style="color:#991B1B;margin:0;">{len(missing)}</h2>
        <p style="margin:0;">Missing Skills</p>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)
        
        # Skill Comparison Chart
        left, right = st.columns([3, 1])

        with left:
            st.markdown("###  Skill Match Overview")

            fig = go.Figure()
            fig.add_bar(
                x=df_skills["Skill"],
                y=df_skills["Resume Skill %"],
                name="Resume Skills",
                marker_color="#3B82F6"
            )
            fig.add_bar(
                x=df_skills["Skill"],
                y=df_skills["Job Requirement %"],
                name="Job Requirements",
                marker_color="#10B981"
            )

            fig.update_layout(
                barmode="group",
                height=350,
                yaxis_title="Percentage (%)",
                xaxis_tickangle=45
            )

            st.plotly_chart(fig, use_container_width=True)
 
            st.markdown("###  Upskilling Recommendations")
            st.info(
    "ℹ️ Recommendations are based on skills that are **missing** or **partially matched** compared to the job requirements."
)
            if missing or partial:
                for skill in sorted(missing | partial):
                    st.warning(
                        f"Improve **{skill.title()}** through courses and hands-on projects"
                    )
            else:
                st.success("Perfect match! All required skills are present.")





            # ================= IMPROVEMENT PRIORITY MATRIX =================
                     
        # ================= END REAL DATA MATRIX =================


        
        # ================= END HIRING READINESS OVERVIEW =================

# =======================
# 📈 Skill Similarity Scores - NEW DESIGN
# =======================

st.markdown('<div class="similarity-section">', unsafe_allow_html=True)
st.markdown('<div class="similarity-title">📊 Skill Similarity Scores</div>', unsafe_allow_html=True)
st.markdown('<div class="colorful-caption"><em>How closely your skills align with the job requirements</em></div>', unsafe_allow_html=True)

# Create a more structured display with better formatting
skills_display = []
for skill, score in skill_scores.items():
    percentage = score * 100
    # Better categorization
    if percentage >= 70:
        match_level = "Strong Match"
        color = "#28a745"  # Green
    elif percentage >= 40:
        match_level = "Moderate Match"
        color = "#ffc107"  # Yellow/Orange
    elif percentage >= 10:
        match_level = "Basic Match"
        color = "#fd7e14"  # Orange
    else:
        match_level = "Low Match"
        color = "#dc3545"  # Red
    
    skills_display.append({
        "skill": skill,
        "percentage": percentage,
        "score": score,
        "match_level": match_level,
        "color": color
    })

# Sort by percentage (highest first)
skills_display.sort(key=lambda x: x["percentage"], reverse=True)

# Create metrics row
col1, col2, col3, col4 = st.columns(4)
with col1:
    avg_match = sum([s["percentage"] for s in skills_display]) / len(skills_display) if skills_display else 0
    st.metric("Average Match", f"{avg_match:.1f}%")
with col2:
    strong_matches = len([s for s in skills_display if s["percentage"] >= 70])
    st.metric("Strong Matches", strong_matches)
with col3:
    moderate_matches = len([s for s in skills_display if s["percentage"] >= 40 and s["percentage"] < 70])
    st.metric("Moderate Matches", moderate_matches)
with col4:
    low_matches = len([s for s in skills_display if s["percentage"] < 40])
    st.metric("Low Matches", low_matches)

# Display skills in a cleaner table format
st.markdown("---")
st.markdown("#### Skill Match Details")

# Create a DataFrame for better display
if skills_display:
    df_skills = pd.DataFrame({
        "Skill": [s["skill"] for s in skills_display],
        "Match %": [f"{s['percentage']:.1f}%" for s in skills_display],
        "Level": [s["match_level"] for s in skills_display]
    })
    
    # Apply color coding
    def color_rows(val):
        if "Strong" in val:
            return 'background-color: #d4edda; color: #155724;'
        elif "Moderate" in val:
            return 'background-color: #fff3cd; color: #856404;'
        elif "Basic" in val:
            return 'background-color: #ffe5d0; color: #993d00;'
        else:
            return 'background-color: #f8d7da; color: #721c24;'
    
    styled_df = df_skills.style.applymap(color_rows, subset=['Level'])
    st.dataframe(styled_df, use_container_width=True, height=400)

# Create the horizontal bar chart
st.markdown("---")
st.markdown("#### Skill Similarity Visualization")

if skills_display:
    skills = [s["skill"] for s in skills_display]
    percentages = [s["percentage"] for s in skills_display]
    colors = [s["color"] for s in skills_display]

    # ================= ENHANCED VISUALIZATION =================
    st.markdown("### 📊 Skill Similarity Visualization")

    if skills_display:
        skills = [s["skill"] for s in skills_display]
        percentages = [s["percentage"] for s in skills_display]
        colors = [s["color"] for s in skills_display]
        
        # Create an enhanced horizontal bar chart
        fig = go.Figure()
        
        # Add bars with gradient effect
        fig.add_trace(go.Bar(
            x=percentages,
            y=skills,
            orientation='h',
            marker=dict(
                color=colors,
                line=dict(color='rgba(0,0,0,0.3)', width=1),
                opacity=0.85
            ),
            text=[f"{p:.1f}%" for p in percentages],
            textposition='outside',
            textfont=dict(size=12, color='#333'),
            hoverinfo='text',
            hovertext=[f"<b>{skill}</b><br>Match: {p:.1f}%<br>Level: {match_level}" 
                    for skill, p, match_level in zip(skills, percentages, [s["match_level"] for s in skills_display])]
        ))
        
        # Update layout for better visualization
        fig.update_layout(
            height=max(350, len(skills) * 35),
            xaxis=dict(
                title="Similarity (%)",
                range=[0, 100],
                ticksuffix="%",
                title_font=dict(size=14, color="#2c3e50", family="Arial Black"),
                tickfont=dict(size=12),
                gridcolor='rgba(0,0,0,0.1)',
                gridwidth=1
            ),
            yaxis=dict(
                title="Skills",
                tickfont=dict(size=13),
                automargin=True,
                categoryorder='total ascending'
            ),
            margin=dict(l=150, r=40, t=50, b=50),
            plot_bgcolor="rgba(248, 249, 250, 0.8)",
            paper_bgcolor="white",
            showlegend=False,
            title=dict(
                text="📈 Skill Match Percentage Visualization",
                font=dict(size=18, color="#2c3e50", family="Arial Black"),
                x=0.5,
                xanchor="center",
                y=0.95
            )
        )
        
        # Add reference zones with better styling
        fig.add_vrect(x0=70, x1=100, fillcolor="green", opacity=0.1, layer="below", line_width=0)
        fig.add_vrect(x0=40, x1=70, fillcolor="orange", opacity=0.1, layer="below", line_width=0)
        fig.add_vrect(x0=10, x1=40, fillcolor="red", opacity=0.1, layer="below", line_width=0)
        
        # Add reference lines with labels
        fig.add_vline(x=70, line_dash="dash", line_color="green", opacity=0.7, 
                    annotation_text="Strong (70%+)", annotation_position="top")
        fig.add_vline(x=40, line_dash="dash", line_color="orange", opacity=0.7,
                    annotation_text="Moderate (40%+)", annotation_position="top")
        fig.add_vline(x=10, line_dash="dash", line_color="red", opacity=0.7,
                    annotation_text="Basic (10%+)", annotation_position="top")
        
        st.plotly_chart(fig, use_container_width=True)

    # ================= MATCH LEVEL LEGEND =================

    # ================= RECOMMENDATION BASED ON SCORES =================
    low_skills = [s for s in skills_display if s["percentage"] < 40]
    if low_skills:
        st.markdown("""
        <div style="
            background: linear-gradient(135deg, #fff5f5 0%, #fed7d7 100%);
            padding: 20px;
            border-radius: 10px;
            border: 2px dashed #fc8181;
            margin: 20px 0;
        ">
            <h4 style="color: #c53030; margin-top: 0;">🚨 Areas Needing Improvement</h4>
            <p style="color: #742a2a;">
                <strong>Focus on these skills first:</strong> 
                {skills_list}
            </p>
            <div style="
                background: white;
                padding: 15px;
                border-radius: 8px;
                margin-top: 10px;
                border-left: 4px solid #4299e1;
            ">
                <strong>💡 Recommendation:</strong> Consider taking online courses or 
                working on projects to improve these skills. Even small improvements 
                (from 0% to 30%) can significantly increase your match score.
            </div>
        </div>
        """.format(skills_list=", ".join([s["skill"] for s in low_skills[:3]])), unsafe_allow_html=True)




        with right:
            st.markdown("###  Role View")
            selected_role = st.radio("", ["Job Seeker", "Recruiter"], horizontal=True, key="role")
            
            if len(all_skills) >= 3:
                radar = go.Figure()

                # Decide axis labels and corresponding scores based on view
                if selected_role == "Job Seeker":
                    theta_labels = [s.title() for s in all_skills] + [all_skills[0].title()]  # All unique skills
                    jd_r = jd_scores + [jd_scores[0]]                                            # Purple = JD requirement
                    resume_r = resume_scores + [resume_scores[0]]                                # Blue = Your resume coverage
                    partial_r = [55 if s.lower() in partial else 0 for s in all_skills] + [0]
                else:
                    # Recruiter: ONLY JD skills on axis
                    jd_only_skills = [s for s in all_skills if s.lower() in jd_skill_names]
                    if not jd_only_skills:
                        jd_only_skills = all_skills[:1]  # fallback if no JD skills

                    theta_labels = [s.title() for s in jd_only_skills] + [jd_only_skills[0].title()]

                    # Map JD scores to JD-only axis
                    jd_r = [jd_scores[all_skills.index(s)] if s in all_skills else 0 for s in jd_only_skills] + [0]

                    # Matched / Partial only on JD skills
                    matched_r = [85 if s.lower() in matched else 0 for s in jd_only_skills] + [0]
                    partial_r = [55 if s.lower() in partial else 0 for s in jd_only_skills] + [0]

                    # No resume line in recruiter view (gaps clear)

                # JD base (always strong green)
                radar.add_trace(go.Scatterpolar(
                    r=jd_r,
                    theta=theta_labels,
                    fill="toself",
                    name="Job Requirement",
                    line_color="#A855F7",          # <-- Purple (reddish-purple)
                    fillcolor="rgba(168, 85, 247, 0.25)",   # lighter fill
                    opacity=0.85
                ))

                if selected_role == "Job Seeker":
                    # Strong resume profile for job seeker
                    radar.add_trace(go.Scatterpolar(
                        r=resume_r,
                        theta=theta_labels,
                        fill="toself",
                        name="Your Full Profile (Resume Skills)",
                        line_color="#3B82F6",          # <-- Blue
                        fillcolor="rgba(59, 130, 246, 0.45)",   # a bit stronger fill
                        opacity=0.95
                    ))
                    title_text = "Your Complete Skill Alignment vs Job Requirement"

                else:  # Recruiter: Only matched + partial overlays
                    # Matched (darker green)
                    radar.add_trace(go.Scatterpolar(
                        r=matched_r,
                        theta=theta_labels,
                        fill="toself",
                        name="Matched Skills",
                        line_color="#22C55E",
                        fillcolor="rgba(34, 197, 94, 0.6)",
                        opacity=0.9
                    ))

                    # Partial (yellow)
                    radar.add_trace(go.Scatterpolar(
                        r=partial_r,
                        theta=theta_labels,
                        fill="toself",
                        name="Partial Matches",
                        line_color="#F59E0B",
                        fillcolor="rgba(245, 158, 11, 0.6)",
                        opacity=0.8
                    ))

                    # Optional: Red outline for missing (uncomment if you want red border on missing)
                    missing_r = [jd_scores[all_skills.index(s)] if s.lower() not in matched and s.lower() not in partial else 0 for s in jd_only_skills] + [0]
                    radar.add_trace(go.Scatterpolar(
                       r=missing_r,
                        theta=theta_labels,
                        mode="lines",
                        name="Missing Skills (Gaps)",
                        line=dict(color="#EF4444", width=3, dash="dot"),
                        fill=None,
                        showlegend=True
                    ))

                    title_text = "Candidate Fit vs Job Requirement (Recruiter View - Gaps Highlighted)"

                radar.update_layout(
                    polar=dict(
                        radialaxis=dict(range=[0, 100], visible=True, tickfont=dict(size=12)),
                        angularaxis=dict(
                            showticklabels=True,
                            tickfont=dict(size=11),           # Smaller font if many skills
                            rotation=90,                      # Rotate labels
                            direction="clockwise"
                        )
                    ),
                    height=520,                               # Bigger chart
                    showlegend=True,
                    legend=dict(orientation="h", yanchor="bottom", y=-0.4, xanchor="center", x=0.5),
                    title=dict(text="Your Complete Skill Alignment vs Job Requirement", x=0.5, xanchor="center", font=dict(size=18)),
                    margin=dict(t=120, b=220, l=80, r=80)     # Extra bottom space for labels
                )

                st.plotly_chart(radar, use_container_width=True)
            else:
                st.info("Need at least 3 skills for radar chart")
                
  # ================= NEW INNOVATIVE FEATURE =================
       
        st.markdown('<div class="priority-section">', unsafe_allow_html=True)
        st.markdown('<div class="priority-title"> Skill Priority Radar</div>', unsafe_allow_html=True)

        st.markdown('<div class="colorful-caption">Ranks missing skills based on how critical they are in the job description.</div>', unsafe_allow_html=True)

        if len(missing) == 0:
            st.markdown('<div class="colorful-info-box">🎉 No missing skills — you already meet all job requirements.</div>', unsafe_allow_html=True)
        else:
            skill_priority = {}

            for skill in missing:
                pattern = r"\b" + re.escape(skill.lower()) + r"\b"
                frequency = len(re.findall(pattern, jd_text.lower()))
                skill_priority[skill] = max(1, frequency)

            priority_df = pd.DataFrame({
                "Skill": skill_priority.keys(),
                "Importance Score": skill_priority.values()
            }).sort_values("Importance Score", ascending=False)

            st.markdown('<div class="high-impact-section">', unsafe_allow_html=True)
            st.markdown('<div class="high-impact-title">🔥 High-Impact Skills to Focus On First</div>', unsafe_allow_html=True)
            
            # Display each missing skill with its importance
            for idx, (skill, score) in enumerate(priority_df.values):
                st.markdown(f"""
                <div class="missing-skill-item">
                    <strong>{idx+1}. {skill.title()}</strong> - Importance: <span style="color: #E91E63; font-weight: bold;">{score}</span>
                </div>
                """, unsafe_allow_html=True)
            
            st.markdown('</div>', unsafe_allow_html=True)

            

            top_skill = priority_df.iloc[0]["Skill"]
            
            st.markdown(f"""
            <div class="start-with-box">
                 <strong>Start with "{top_skill.title()}"</strong> — it appears most frequently 
                in the job description and has the highest hiring impact.
            </div>
            """, unsafe_allow_html=True)

        st.markdown('</div>', unsafe_allow_html=True)
        # ================= END NEW FEATURE =================  


        # Export Section - FIXED
        st.markdown("---")
        st.markdown("### ▼ Export Reports")
        
        def create_skill_charts(matched, partial, missing, overall_match):
            # ---- Skill Distribution Pie Chart ----
            labels = ["Matched", "Partial", "Missing"]
            sizes = [len(matched), len(partial), len(missing)]

            plt.figure(figsize=(5, 5))   # bigger + clean
            plt.pie(
                sizes,
                labels=labels,
                autopct="%1.1f%%",
                startangle=90
            )
            plt.title("Skill Distribution")
            plt.tight_layout()
            plt.savefig("skills_distribution.png", dpi=200)
            plt.close()   # VERY IMPORTANT


            # ---- Overall Match Bar Chart ----
            plt.figure(figsize=(8, 3))
            plt.bar(["Overall Match"], [overall_match])
            plt.ylim(0, 100)
            plt.title("Overall Match Rate (%)")
            plt.tight_layout()
            plt.savefig("overall_match.png", dpi=200)
            plt.close()   # 🔥 THIS IS CRITICAL





        def generate_pdf_report():
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()

            generated_on = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            candidate_name, linkedin_url = extract_name_and_linkedin(resume_text)

            # ================= HEADER =================
            pdf.set_text_color(33, 150, 243)
            pdf.set_font("Arial", "B", 18)
            pdf.cell(0, 12, "Skill Gap Analysis Report", ln=True, align="C")
            pdf.ln(8)
            pdf.set_text_color(0, 0, 0)

            # ================= CANDIDATE INFO =================
            pdf.set_font("Arial", "", 11)
            pdf.cell(0, 8, f"Candidate Name: {candidate_name}", ln=True)
            if linkedin_url:
                pdf.cell(0, 8, f"LinkedIn: {linkedin_url}", ln=True)
            pdf.ln(6)

            # ================= OVERALL MATCH =================
            pdf.set_font("Arial", "B", 13)
            if overall_match >= 70:
                pdf.set_text_color(56, 142, 60)
            elif overall_match >= 50:
                pdf.set_text_color(245, 124, 0)
            else:
                pdf.set_text_color(211, 47, 47)

            pdf.cell(0, 10, f"Overall Match Rate: {overall_match}%", ln=True)
            pdf.set_text_color(0, 0, 0)
            pdf.ln(6)

            # ================= SKILL SUMMARY =================
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 8, "Skill Summary", ln=True)

            pdf.set_font("Arial", "", 11)
            pdf.set_text_color(46, 125, 50)
            pdf.cell(0, 7, f"Matched Skills: {len(matched)}", ln=True)

            pdf.set_text_color(245, 124, 0)
            pdf.cell(0, 7, f"Partially Matched Skills: {len(partial)}", ln=True)

            pdf.set_text_color(198, 40, 40)
            pdf.cell(0, 7, f"Missing Skills: {len(missing)}", ln=True)

            pdf.set_text_color(0, 0, 0)
            pdf.ln(8)

            # ================= GRAPHS =================
            create_skill_charts(matched, partial, missing, overall_match)

            # Overall Match
# ----- Overall Match Bar Chart -----
# ----- Overall Match Bar Chart -----
            if pdf.get_y() > 170:
                pdf.add_page()

            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 8, "Overall Match Rate", ln=True)
            pdf.ln(3)

            pdf.image("overall_match.png", x=30, w=150)
            pdf.ln(12)


            # ----- Skills Distribution Pie Chart -----
# ----- Skills Distribution Pie Chart -----
            if pdf.get_y() > 160:
                pdf.add_page()

            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 8, "Skills Distribution", ln=True)
            pdf.ln(3)

            pdf.image("skills_distribution.png", x=45, w=120)
            pdf.ln(10)




 # <-- THIS LINE IS CRITICAL


            # ================= TECHNICAL SKILLS =================
            pdf.ln(10)
            pdf.set_font("Arial", "B", 12)
            pdf.set_text_color(198, 40, 40)
            pdf.cell(0, 8, "Missing Technical Skills", ln=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", "", 11)

            tech_found = False
            for skill in sorted(missing):
                if any(s["name"].lower() == skill and s["type"] == "Technical" for s in jd_skills):
                    pdf.cell(0, 7, f"- {skill.title()}", ln=True)
                    tech_found = True

            if not tech_found:
                pdf.cell(0, 7, "None", ln=True)

            pdf.ln(6)

            # ================= SOFT SKILLS =================
            pdf.set_font("Arial", "B", 12)
            pdf.set_text_color(245, 124, 0)
            pdf.cell(0, 8, "Missing Soft Skills", ln=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", "", 11)

            soft_found = False
            for skill in sorted(missing):
                if any(s["name"].lower() == skill and s["type"] == "Soft" for s in jd_skills):
                    pdf.cell(0, 7, f"- {skill.title()}", ln=True)
                    soft_found = True

            if not soft_found:
                pdf.cell(0, 7, "None", ln=True)

            pdf.ln(8)

            # ================= RECOMMENDATIONS =================
            pdf.set_font("Arial", "B", 13)
            pdf.set_text_color(30, 136, 229)
            pdf.cell(0, 9, "Recommendations", ln=True)
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("Arial", "", 11)

            if overall_match >= 70:
                pdf.multi_cell(
                    0, 7,
                    "You have a strong skill match for this role. Maintain your current skills "
                    "and continue updating your knowledge to stay competitive."
                )
            elif overall_match >= 50:
                pdf.multi_cell(
                    0, 7,
                    "You partially match the job requirements. Focus on improving missing "
                    "and partially matched skills through projects and targeted learning."
                )
            else:
                pdf.multi_cell(
                    0, 7,
                    "There is a significant skill gap. Consider structured learning, "
                    "certifications, and hands-on practice before applying to similar roles."
                )

            # ================= FOOTER (DATE AT BOTTOM) =================
            pdf.ln(12)
            pdf.set_font("Arial", "I", 9)
            pdf.set_text_color(120, 120, 120)
            pdf.cell(0, 8, f"Generated on: {generated_on}", align="C")
            pdf.set_text_color(0, 0, 0)

            # Cleanup chart images
            if os.path.exists("skills_distribution.png"):
                os.remove("skills_distribution.png")
            if os.path.exists("overall_match.png"):
                os.remove("overall_match.png")

            return pdf.output(dest="S").encode("latin-1", "ignore")

        
        def generate_csv_report():
            """Generate CSV report with skill scores instead of Yes/No"""
            import io
            output = io.StringIO()

            # CSV Header
            output.write("Skill,Status,Resume Score (%),Job Requirement Score (%)\n")

            for skill in sorted(all_skills):
                if skill in matched:
                    status = "Matched"
                    resume_score = 85
                    jd_score = 90
                elif skill in partial:
                    status = "Partial"
                    resume_score = 55
                    jd_score = 80
                else:
                    status = "Missing"
                    resume_score = 20
                    jd_score = 90

                output.write(
                    f"{skill.title()},{status},{resume_score},{jd_score}\n"
                )

            return output.getvalue().encode("utf-8")

        
        # Download buttons
        st.markdown('<div class="sub-section">Export Reports</div>', unsafe_allow_html=True)
        col1, col2 = st.columns(2)

        with col1:
            try:
                pdf_data = generate_pdf_report()
                st.download_button(
                    label="Download PDF Report",
                    data=pdf_data,
                    file_name="skill_gap_report.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
            except Exception as e:
                st.markdown('<div class="status-error">PDF generation failed</div>', unsafe_allow_html=True)

        with col2:
            try:
                csv_data = generate_csv_report()
                st.download_button(
                    label="Download CSV Report",
                    data=csv_data,
                    file_name="skill_gap_data.csv",
                    mime="text/csv",
                    use_container_width=True
                )
            except Exception as e:
                st.markdown('<div class="status-error">CSV generation failed</div>', unsafe_allow_html=True)
                
 # ================= FINAL COMPLETION MESSAGE =================
        st.markdown("---")
        st.success("✅ **Completed: Dashboard visualization and report generation completed successfully.**")
        
                
                

    else:
        st.warning("No skills detected for visualization")

else:
    st.info("👈 Please upload both a resume and a job description to start the analysis.")
    
    
